"""
Routes and views for the flask application.
"""
import os
import time
import base64
from datetime import datetime
from flask import render_template, request, redirect, url_for, send_from_directory, make_response, session, flash, g
from SL import app, UPLOAD_FOLDER
from SL.helper.helper import allowed_file
from SL.functions.ht_space import ht_space_inspect
from SL.functions.ht_name import ht_name_inspect
from SL.functions.gys_name import gys_name_inspect, str_name_inspect
from SL.functions.num_case import num_case_inspect
from SL.functions.pay_conditions import pay_inspect
from SL.functions.clause import clause_inspect
import docx
from SL.forms.ht_name_form import HtNameForm
from SL.forms.gys_name_form import GysNameForm
from SL.forms.pay_conditions_form import PayConditionsForm
from SL.forms.clause_form import ClauseForm


@app.route("/all", methods=['GET'])
def process_all():
    pass


@app.route('/')
@app.route('/index')
def index():
    """Renders the index page."""
    #return render_template(
    #    'index.html',
    #    title='演示系统',
    #    year=datetime.now().year)
    return render_template('index.html',title='',year = datetime.now().year)

@app.route('/clause',methods=['POST','GET'])
def clause():
    new_filename = None
    form = ClauseForm()
    #filename = session['filename']
    res = None
    if form.validate_on_submit():
        if session.get('filename',None) is None:
            flash('未上传文档')
            filename= None
        else:
            filename = session.get("filename",None)
            res = clause_inspect(form.clauses.data,docx.Document(os.path.join(UPLOAD_FOLDER,filename)))
            #print(res)
            #new_filename = bytes.decode(base64.b64encode(bytes(str(time.time()),encoding='utf-8'))) +'_'+ filename.split('_')[1]
            #doc.save(os.path.join(UPLOAD_FOLDER,new_filename))
    return render_template(
        'clause.html',
        title='合同条款',
        form=form,
        filename=None,
        res = {} if res is None else res
    )

@app.route('/gys_name', methods=["POST","GET"])
def gys_name():
    new_filename = None
    form = GysNameForm()
    if form.validate_on_submit():
        if session.get('filename',None) is None:
            flash('未上传文档')
            filename= None
        else:
            filename = session['filename']
            doc = gys_name_inspect(form.gys_name.data,docx.Document(os.path.join(UPLOAD_FOLDER,filename)))
            doc = str_name_inspect(form.str_name.data, doc)
            new_filename = bytes.decode(base64.b64encode(bytes(str(time.time()),encoding='utf-8'))) +'_'+ filename.split('_')[1]
            doc.save(os.path.join(UPLOAD_FOLDER,new_filename))
    return render_template(
        'gys_name.html',
        title='合同主体',
        form=form,
        filename=new_filename
    )

@app.route('/ht_name', methods=['GET','POST'])
def ht_name():
    new_filename = None
    form = HtNameForm()
    if form.validate_on_submit():
        if session.get('filename',None) is None:
            flash('未上传文档')
            filename= None
        else:
            filename = session['filename']
            doc = ht_name_inspect(form.name.data,docx.Document(os.path.join(UPLOAD_FOLDER,filename)))
            new_filename = bytes.decode(base64.b64encode(bytes(str(time.time()),encoding='utf-8'))) +'_'+ filename.split('_')[1]
            doc.save(os.path.join(UPLOAD_FOLDER,new_filename))

    return render_template(
        'ht_name.html',
        title='合同名称',
        filename = new_filename,
        form=form
    )

@app.route('/ht_space',methods=["POST","GET"])
def ht_space():
    new_filename = None
    if session.get('filename',None) is None:
        flash('未上传文档')
        filename = None
    else:
        filename = session['filename']
        doc = ht_space_inspect(docx.Document(os.path.join(UPLOAD_FOLDER,filename)))
        new_filename = bytes.decode(base64.b64encode(bytes(str(time.time()),encoding='utf-8'))) +'_'+ filename.split('_')[1]
        doc.save(os.path.join(UPLOAD_FOLDER,new_filename))
    return render_template(
        'ht_space.html',
        title='合同空白',
        filename=new_filename
    )

@app.route('/num_case', methods=["POST","GET"])
def num_case():
    new_filename = None
    if session.get('filename',None) is None:
        flash('未上传文档')
        filename = None
    else:
        filename = session['filename']
        doc = num_case_inspect(docx.Document(os.path.join(UPLOAD_FOLDER,filename)))
        new_filename = bytes.decode(base64.b64encode(bytes(str(time.time()),encoding='utf-8'))) +'_'+ filename.split('_')[1]
        doc.save(os.path.join(UPLOAD_FOLDER,new_filename))
    return render_template(
        'num_case.html',
        title='金额大小写',
        filename=new_filename
    )

@app.route('/pay_conditions', methods=['POST','GET'])
def pay_conditions():
    new_filename = None
    form = PayConditionsForm()
    if form.validate_on_submit():
        if session.get('filename',None) is None:
            flash('未上传文档')
            filename= None
        else:
            filename = session['filename']
            doc = pay_inspect(form.conditions.data,docx.Document(os.path.join(UPLOAD_FOLDER,filename)))
            new_filename = bytes.decode(base64.b64encode(bytes(str(time.time()),encoding='utf-8'))) +'_'+ filename.split('_')[1]
            doc.save(os.path.join(UPLOAD_FOLDER,new_filename))


    return render_template(
        'pay_conditions.html',
        title='支付条件',
        form=form,
        filename=new_filename
    )

@app.route('/upload_word', methods=['POST'])
def upload_word():
    if request.method == 'POST':
        file = request.files['file']
        print(file.filename)
        if file and allowed_file(file.filename):
            filename =  bytes.decode(base64.b64encode(bytes(str(time.time()),encoding='utf-8'))) + '_' + file.filename
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            print(filename)
            session['filename'] = filename
        else:
            flash("文件格式不正确")
    return redirect(url_for('index'))

@app.route('/download/<string:filename>',methods=['GET'])
def download(filename):
    response = make_response(send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True))
    response.headers["Content-Disposition"] = "attachment; filename={}".format(filename.encode().decode('utf-8'))
    return response
