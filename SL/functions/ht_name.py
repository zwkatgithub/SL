import re
import jieba
import docx


def get_all(doc):
    return ''.join([p.text for p in doc.paragraphs])
def get_runs_by_range(doc, start, end):
    before_start = 0
    ps = []
    i = 0
    x_start = 0
    while i < len(doc.paragraphs):
        if before_start + len(doc.paragraphs[i].text) >= start:
            x_start = before_start
            while before_start <= end:
                ps.append(doc.paragraphs[i])
                before_start += len(doc.paragraphs[i].text)
                i += 1
            break
        before_start += len(doc.paragraphs[i].text)
        i += 1
    new_start = start - x_start
    new_end = end - x_start
    
    i = 0
    runs = []
    for p in ps:
        runs += p.runs
    before_start = 0
    rs = []
    while i < len(runs):
        if before_start + len(runs[i].text) >= new_start:
            while before_start <= new_end:
                rs.append(runs[i])
                before_start += len(runs[i].text)
                i += 1
            break
        before_start += len(runs[i].text)
        i += 1
    return rs

        
def ht_name_inspect(ht_name ,doc):
    words = jieba.cut(ht_name)
    r = re.compile(''.join(['('+word+')?' for word in words]))
    doc_text = get_all(doc)
    groups = re.findall(r, doc_text)
    ranges = list(set([g for g in [''.join(group) for group in groups] if len(g) > (len(ht_name)//2)]))
    for ht_name_ in ranges:
        if ht_name_ != ht_name:
            for str_ in re.finditer(ht_name_, doc_text):
                runs = get_runs_by_range(doc, str_.start(), str_.end())
                for r in runs:
                    r.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

    return doc

if __name__ == "__main__":
    doc = docx.Document(r"C:\Users\86150\Desktop\1.docx")
    doc = ht_name_inspect('XXXX银行2018年核心系统建设工程之XXX系统配合改造项目', doc)
    doc.save(r"C:\Users\86150\Desktop\2.docx")